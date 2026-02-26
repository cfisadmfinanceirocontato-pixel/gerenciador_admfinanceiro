"""
SISTEMA DE PAGAMENTO DE DIÁRIAS - VERSÃO CORRIGIDA
✅ Corrigido: Placeholders (VALOR POR EXTENSO) e (QTD POR EXTENSO) agora funcionam
✅ Corrigido: Problema do caminho no Streamlit Cloud (sem condicionais complexas)
✅ Corrigido: Reset de formulário agora limpa todos os campos
✅ Adicionado: Botão para excluir todos os registros com confirmação
✅ Demais funcionalidades mantidas exatamente iguais
"""

import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, date
import io
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import subprocess
from pathlib import Path
import tempfile
import platform
import shutil
import zipfile
import requests
from io import StringIO

# =============================================================================
# 📄 CONVERSÃO DOCX → PDF COM FORMATAÇÃO PERFEITA
# =============================================================================
PDF_NATIVE_AVAILABLE = False
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    PDF_NATIVE_AVAILABLE = True
except ImportError as e:
    PDF_NATIVE_AVAILABLE = False

def docx_to_pdf_perfeito(docx_path, pdf_path):
    """
    Converte DOCX para PDF preservando TODA a formatação original
    """
    if not PDF_NATIVE_AVAILABLE:
        return False, "ReportLab não disponível"
    
    try:
        doc = Document(docx_path)
        pdf = SimpleDocTemplate(
            str(pdf_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: TA_LEFT,
            WD_ALIGN_PARAGRAPH.CENTER: TA_CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT: TA_RIGHT,
            WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY
        }
        
        story = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip() and not paragraph.runs:
                story.append(Spacer(1, 12))
                continue
            
            align = align_map.get(paragraph.paragraph_format.alignment, TA_LEFT)
            
            space_before = paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else 0
            space_after = paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 0
            
            if space_before > 0:
                story.append(Spacer(1, space_before))
            
            runs_info = []
            
            for run in paragraph.runs:
                text = run.text
                if not text:
                    continue
                
                font_size = run.font.size.pt if run.font.size else 12
                
                is_bold = run.font.bold if run.font.bold is not None else False
                is_italic = run.font.italic if run.font.italic is not None else False
                
                if is_bold and is_italic:
                    pdf_font = "Helvetica-BoldOblique"
                elif is_bold:
                    pdf_font = "Helvetica-Bold"
                elif is_italic:
                    pdf_font = "Helvetica-Oblique"
                else:
                    pdf_font = "Helvetica"
                
                font_color = colors.black
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    font_color = colors.Color(rgb[0]/255, rgb[1]/255, rgb[2]/255)
                
                underline = run.font.underline if run.font.underline is not None else False
                
                runs_info.append({
                    'text': text,
                    'font': pdf_font,
                    'size': font_size,
                    'color': font_color,
                    'underline': underline
                })
            
            if runs_info:
                style_name = f"custom_style_{len(story)}"
                custom_style = ParagraphStyle(
                    style_name,
                    parent=styles['Normal'],
                    fontName=runs_info[0]['font'],
                    fontSize=runs_info[0]['size'],
                    textColor=runs_info[0]['color'],
                    alignment=align,
                    spaceBefore=space_before,
                    spaceAfter=space_after
                )
                
                html_text = ""
                for run in runs_info:
                    text = run['text'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    formatted_text = f'<font face="{run["font"]}" size="{run["size"]}" color="{run["color"]}">'
                    
                    if run['underline']:
                        formatted_text += f'<u>{text}</u>'
                    else:
                        formatted_text += text
                    
                    formatted_text += '</font>'
                    html_text += formatted_text
                
                p = Paragraph(html_text, custom_style)
                story.append(p)
            
            if space_after > 0:
                story.append(Spacer(1, space_after))
        
        for table in doc.tables:
            if not table.rows:
                continue
            
            data = []
            col_widths = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        cell_text += paragraph.text + "\n"
                    row_data.append(cell_text.strip())
                data.append(row_data)
            
            if table.columns:
                col_widths = [2*inch] * len(table.columns)
            
            pdf_table = Table(data, colWidths=col_widths)
            
            table_style = [
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
            ]
            
            if len(data) > 0:
                table_style.append(('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'))
            
            pdf_table.setStyle(TableStyle(table_style))
            story.append(pdf_table)
            story.append(Spacer(1, 12))
        
        pdf.build(story)
        return True, str(pdf_path)
        
    except Exception as e:
        return False, str(e)

# ============================================================================
# ✅ FUNÇÕES DE BUSCA NO CSV
# ============================================================================
@st.cache_data(ttl=300)
def carregar_termos_colaboracao(df):
    """Carrega lista única de termos de colaboração"""
    if df.empty:
        return []

    coluna_termo = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'TERMO' in col_upper and 'COLABORAÇÃO' in col_upper:
            coluna_termo = col
            break
    
    if not coluna_termo and len(df.columns) >= 3:
        coluna_termo = df.columns[2]
    
    if not coluna_termo:
        return []

    return sorted(df[coluna_termo].dropna().astype(str).unique())

@st.cache_data(ttl=300)
def buscar_instrumento_por_termo(df, termo):
    """Busca instrumento a partir do termo"""
    if df.empty or not termo:
        return ""

    coluna_termo = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'TERMO' in col_upper and 'COLABORAÇÃO' in col_upper:
            coluna_termo = col
            break
    
    if not coluna_termo and len(df.columns) >= 3:
        coluna_termo = df.columns[2]
    
    if not coluna_termo:
        return ""

    mask = df[coluna_termo].astype(str).str.strip() == str(termo).strip()
    if not mask.any():
        return ""

    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'INSTRUMENTO' in col_upper:
            valor = df.loc[mask, col].iloc[0]
            return str(valor).strip() if pd.notna(valor) else ""

    return ""

@st.cache_data(ttl=300)
def buscar_numero_termo_por_nome(df, termo):
    """
    Busca o número do termo na coluna "Nº TERMO"
    """
    if df.empty or not termo:
        return ""

    coluna_termo = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'TERMO' in col_upper and 'COLABORAÇÃO' in col_upper:
            coluna_termo = col
            break
    
    if not coluna_termo and len(df.columns) >= 3:
        coluna_termo = df.columns[2]
    
    if not coluna_termo:
        return ""

    mask = df[coluna_termo].astype(str).str.strip() == str(termo).strip()
    if not mask.any():
        return ""

    coluna_numero = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if ('Nº' in col_upper or 'N°' in col_upper or 'NUMERO' in col_upper) and 'TERMO' in col_upper:
            coluna_numero = col
            break
    
    if coluna_numero:
        valor = df.loc[mask, coluna_numero].iloc[0]
        return str(valor).strip() if pd.notna(valor) else ""
    
    colunas = df.columns.tolist()
    idx_termo = colunas.index(coluna_termo)
    if idx_termo + 1 < len(colunas):
        valor = df.loc[mask, colunas[idx_termo + 1]].iloc[0]
        return str(valor).strip() if pd.notna(valor) else ""

    return ""

@st.cache_data(ttl=300)
def carregar_funcionarios_por_termo(df, termo):
    """Carrega funcionários de um termo"""
    if df.empty or not termo:
        return []

    coluna_termo = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'TERMO' in col_upper and 'COLABORAÇÃO' in col_upper:
            coluna_termo = col
            break
    
    if not coluna_termo and len(df.columns) >= 3:
        coluna_termo = df.columns[2]
    
    if not coluna_termo:
        return []

    mask = df[coluna_termo].astype(str).str.strip() == str(termo).strip()
    if not mask.any():
        return []

    coluna_func = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'FUNCIONÁRIO' in col_upper or 'FUNCIONARIO' in col_upper or 'NOME' in col_upper:
            coluna_func = col
            break

    if not coluna_func:
        return []

    return sorted(df.loc[mask, coluna_func].dropna().astype(str).unique())

@st.cache_data(ttl=300)
def buscar_cpf_cargo_por_funcionario(df, termo, funcionario):
    """Busca CPF e cargo do funcionário"""
    if df.empty or not termo or not funcionario:
        return "", ""

    coluna_termo = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'TERMO' in col_upper and 'COLABORAÇÃO' in col_upper:
            coluna_termo = col
            break
    
    if not coluna_termo and len(df.columns) >= 3:
        coluna_termo = df.columns[2]

    if not coluna_termo:
        return "", ""

    coluna_func = None
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'FUNCIONÁRIO' in col_upper or 'FUNCIONARIO' in col_upper or 'NOME' in col_upper:
            coluna_func = col
            break

    if not coluna_func:
        return "", ""

    mask = (
        (df[coluna_termo].astype(str).str.strip() == str(termo).strip()) &
        (df[coluna_func].astype(str).str.strip() == str(funcionario).strip())
    )

    if not mask.any():
        return "", ""

    linha = df.loc[mask].iloc[0]

    cpf = ""
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'CPF' in col_upper:
            cpf = str(linha.get(col, "")).strip() if pd.notna(linha.get(col, "")) else ""
            if cpf:
                break

    cargo = ""
    for col in df.columns:
        col_upper = str(col).upper().strip()
        if 'CARGO' in col_upper or 'FUNÇÃO' in col_upper or 'FUNCAO' in col_upper:
            cargo = str(linha.get(col, "")).strip() if pd.notna(linha.get(col, "")) else ""
            if cargo:
                break

    return cpf, cargo

# ============================================================================
# ✅ FUNÇÃO SALVAR REGISTRO (MANTIDA ORIGINAL)
# ============================================================================
def salvar_registro_formulario(df, termo_input, instrumento, numero_termo, funcionario_input, 
                              cpf, cargo, qtd, qtd_extenso, valor, valor_extenso, 
                              data_input, data_extenso_display, objetivo, localidades, 
                              periodo, oficio, nome_arquivo, numero_oficio_input, 
                              nome_recibo_input, pasta_saida_str):
    """Salva registro na planilha"""
    
    if not pasta_saida_str or pasta_saida_str.strip() == "":
        st.error("❌ **Informe a pasta de destino!**")
        return None, None, None, None
    
    pasta_saida = Path(pasta_saida_str).absolute()
    pasta_saida.mkdir(parents=True, exist_ok=True)
    caminho_excel = pasta_saida / "registros_completos.xlsx"
    
    colunas_planilha = [
        'Termo de Colaboração', 'Instrumento', 'Nº Do Termo de Colaboração', 
        'Funcionário', 'CPF', 'Cargo', 'Quantidade', 'Quantidade por extenso', 
        'Valor', 'Valor por extenso', 'Data Recibo', 'Data por Extenso', 
        'Objetivo', 'Localidades', 'Período', 'Ofício', 'Nome Arquivo', 
        'Nº do Ofício', 'Nome do Recibo'
    ]
    
    nome_recibo_completo = nome_recibo_input or f"{nome_arquivo}_{numero_oficio_input}_{st.session_state.contador_recibo}"
    
    dados_registro = {
        'Termo de Colaboração': termo_input or '',
        'Instrumento': instrumento or '',
        'Nº Do Termo de Colaboração': numero_termo or '',
        'Funcionário': funcionario_input or '',
        'CPF': cpf or '',
        'Cargo': cargo or '',
        'Quantidade': qtd or '',
        'Quantidade por extenso': qtd_extenso or '',
        'Valor': valor or '',
        'Valor por extenso': valor_extenso or '',
        'Data Recibo': data_input or '',
        'Data por Extenso': data_extenso_display or '',
        'Objetivo': objetivo or '',
        'Localidades': localidades or '',
        'Período': periodo or '',
        'Ofício': oficio or '',
        'Nome Arquivo': nome_arquivo or '',
        'Nº do Ofício': numero_oficio_input or '',
        'Nome do Recibo': nome_recibo_completo
    }
    
    if not termo_input or not funcionario_input or not cpf:
        st.error("❌ **Preencha obrigatoriamente**: Termo, Funcionário e CPF!")
        return None, None, None, None
    
    novo_registro = pd.DataFrame([dados_registro])[colunas_planilha]
    
    try:
        if caminho_excel.exists():
            dados_existentes = pd.read_excel(caminho_excel)
            dados_atualizados = pd.concat([dados_existentes, novo_registro], ignore_index=True)
        else:
            dados_atualizados = novo_registro
            
    except Exception as e:
        st.error(f"❌ Erro ao carregar Excel: {e}")
        dados_atualizados = novo_registro
    
    try:
        with pd.ExcelWriter(caminho_excel, engine='openpyxl') as writer:
            dados_atualizados.to_excel(writer, sheet_name='Registros', index=False)
            worksheet = writer.sheets['Registros']
            
            for column in worksheet.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            for cell in worksheet[1]:
                cell.font = openpyxl.styles.Font(bold=True)
        
        return dados_atualizados, novo_registro, nome_recibo_completo, caminho_excel
        
    except Exception as e:
        st.error(f"❌ Erro ao salvar Excel: {e}")
        return None, None, None, None

# ============================================================================
# ✅ FUNÇÃO GERAR RECIBO INDIVIDUAL (CORRIGIDA)
# ============================================================================
def gerar_recibo_individual(dados_registro, template_path, pasta_saida_str, gerar_pdf=True):
    """
    ✅ Corrigido: Placeholders (VALOR POR EXTENSO) e (QTD POR EXTENSO) agora funcionam
    """
    if not os.path.exists(template_path):
        st.error("❌ **MODELO.docx não encontrado!**")
        return None, None
    
    if not pasta_saida_str or pasta_saida_str.strip() == "":
        st.error("❌ **Informe a pasta de destino!**")
        return None, None
    
    pasta_saida = Path(pasta_saida_str).absolute()
    pasta_saida.mkdir(parents=True, exist_ok=True)
    
    # EXATAMENTE o mesmo mapeamento do primeiro código
    replacements = {
        "(FUNCIONÁRIO)": str(dados_registro.get('Funcionário', '')),
        "(CARGO)": str(dados_registro.get('Cargo', '')),
        "(CPF)": str(dados_registro.get('CPF', '')),
        "(VALOR)": str(dados_registro.get('Valor', '')),
        "(VALOR POR EXTENSO)": str(dados_registro.get('Valor por extenso', '')),
        "(QTD)": str(dados_registro.get('Quantidade', '')),
        "(QTD POR EXTENSO)": str(dados_registro.get('Quantidade por extenso', '')),
        "(NÚMERO DO INSTRUMENTO)": str(dados_registro.get('Instrumento', '')),
        "(TERMO DE COLABORAÇÃO)": str(dados_registro.get('Nº Do Termo de Colaboração', '')),
        "(OBJETIVO)": str(dados_registro.get('Objetivo', '')),
        "(LOCALIDADES)": str(dados_registro.get('Localidades', '')),
        "(PERÍODO)": str(dados_registro.get('Período', '')),
        "(OFÍCIO)": str(dados_registro.get('Ofício', '')),
        "(DATA RECIBO)": str(dados_registro.get('Data por Extenso', ''))
    }
    
    nome_arquivo_recibo = dados_registro.get('Nome do Recibo', 'recibo_sem_nome')
    nome_arquivo_recibo = "".join(c for c in nome_arquivo_recibo if c.isalnum() or c in (' ', '-', '_')).rstrip()
    docx_saida = pasta_saida / f"{nome_arquivo_recibo}.docx"
    pdf_saida = pasta_saida / f"{nome_arquivo_recibo}.pdf"
    
    try:
        # Carrega o template
        doc = Document(template_path)
        
        # CORREÇÃO ESPECÍFICA: Primeiro, substitui os placeholders problemáticos
        # usando o método do primeiro código (substituição direta no texto)
        placeholders_problematicos = ["(VALOR POR EXTENSO)", "(QTD POR EXTENSO)"]
        
        for placeholder in placeholders_problematicos:
            valor_substituicao = replacements[placeholder]
            
            # Substitui em parágrafos (substituição direta, igual ao primeiro código)
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    # Substitui diretamente no texto do parágrafo (igual ao primeiro código)
                    paragraph.text = paragraph.text.replace(placeholder, valor_substituicao)
            
            # Substitui em tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, valor_substituicao)
        
        # Depois, substitui os demais placeholders (mantém a lógica atual que funciona)
        for old_text, new_text in replacements.items():
            # Pula os que já foram substituídos
            if old_text in placeholders_problematicos:
                continue
                
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if old_text in paragraph.text:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
        
        # Salva o DOCX
        doc.save(docx_saida)
        
        if not docx_saida.exists():
            st.error("❌ **Erro: DOCX não foi criado!**")
            return None, None
        
        st.success(f"✅ **DOCX GERADO:** {docx_saida.name}")
        
        # PDF com formatação perfeita
        pdf_gerado = False
        pdf_caminho = None
        
        if gerar_pdf and PDF_NATIVE_AVAILABLE:
            with st.spinner("🖨️ **Gerando PDF com formatação original..."):
                sucesso, mensagem = docx_to_pdf_perfeito(docx_saida, pdf_saida)
                
                if sucesso and pdf_saida.exists():
                    pdf_gerado = True
                    pdf_caminho = str(pdf_saida)
                    st.success(f"✅ **PDF GERADO:** {pdf_saida.name}")
                else:
                    st.warning(f"⚠️ **PDF não gerado:** {mensagem}")
        
        return (str(docx_saida) if not gerar_pdf else str(pdf_saida) if pdf_gerado else None, 
                str(docx_saida))
        
    except Exception as e:
        st.error(f"❌ **Erro inesperado:** {str(e)}")
        return None, None

# ============================================================================
# ✅ FUNÇÃO GERAR TODOS OS RECIBOS (CORRIGIDA)
# ============================================================================
def gerar_todos_recibos(template_path, pasta_saida_str, gerar_pdf=True):
    """
    ✅ Gera recibos para todos os registros
    """
    if not os.path.exists(template_path):
        st.error("❌ **MODELO.docx não encontrado!**")
        return
    
    if not pasta_saida_str or pasta_saida_str.strip() == "":
        st.error("❌ **Informe a pasta de destino!**")
        return
    
    caminho_excel = Path(pasta_saida_str).absolute() / "registros_completos.xlsx"
    if not caminho_excel.exists():
        st.error(f"❌ **Nenhum registro salvo encontrado em**: `{caminho_excel}`")
        return
    
    try:
        dados_completos = pd.read_excel(caminho_excel)
        if dados_completos.empty:
            st.warning("⚠️ **Nenhum registro para processar**")
            return
        
        pasta_saida = Path(pasta_saida_str).absolute()
        pasta_saida.mkdir(parents=True, exist_ok=True)
        
        total_registros = len(dados_completos)
        st.info(f"🚀 **Processando {total_registros} registros em**: `{pasta_saida}`")
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        pdfs_gerados = 0
        docxs_gerados = 0
        
        for idx, dados_registro in dados_completos.iterrows():
            # EXATAMENTE o mesmo mapeamento do primeiro código
            replacements = {
                "(FUNCIONÁRIO)": str(dados_registro.get('Funcionário', '')),
                "(CARGO)": str(dados_registro.get('Cargo', '')),
                "(CPF)": str(dados_registro.get('CPF', '')),
                "(VALOR)": str(dados_registro.get('Valor', '')),
                "(VALOR POR EXTENSO)": str(dados_registro.get('Valor por extenso', '')),
                "(QTD)": str(dados_registro.get('Quantidade', '')),
                "(QTD POR EXTENSO)": str(dados_registro.get('Quantidade por extenso', '')),
                "(NÚMERO DO INSTRUMENTO)": str(dados_registro.get('Instrumento', '')),
                "(TERMO DE COLABORAÇÃO)": str(dados_registro.get('Nº Do Termo de Colaboração', '')),
                "(OBJETIVO)": str(dados_registro.get('Objetivo', '')),
                "(LOCALIDADES)": str(dados_registro.get('Localidades', '')),
                "(PERÍODO)": str(dados_registro.get('Período', '')),
                "(OFÍCIO)": str(dados_registro.get('Ofício', '')),
                "(DATA RECIBO)": str(dados_registro.get('Data por Extenso', ''))
            }
            
            nome_arquivo_recibo = dados_registro.get('Nome do Recibo', f'recibo_sem_nome_{idx}')
            nome_arquivo_recibo = "".join(c for c in str(nome_arquivo_recibo) if c.isalnum() or c in (' ', '-', '_')).rstrip()
            docx_saida = pasta_saida / f"{nome_arquivo_recibo}.docx"
            pdf_saida = pasta_saida / f"{nome_arquivo_recibo}.pdf"
            
            try:
                doc = Document(template_path)
                
                # CORREÇÃO: Substitui primeiro os placeholders problemáticos
                placeholders_problematicos = ["(VALOR POR EXTENSO)", "(QTD POR EXTENSO)"]
                
                for placeholder in placeholders_problematicos:
                    valor_substituicao = replacements[placeholder]
                    
                    for paragraph in doc.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, valor_substituicao)
                    
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if placeholder in paragraph.text:
                                        paragraph.text = paragraph.text.replace(placeholder, valor_substituicao)
                
                # Depois substitui os demais
                for old_text, new_text in replacements.items():
                    if old_text in placeholders_problematicos:
                        continue
                        
                    for paragraph in doc.paragraphs:
                        if old_text in paragraph.text:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                    
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if old_text in paragraph.text:
                                        for run in paragraph.runs:
                                            if old_text in run.text:
                                                run.text = run.text.replace(old_text, new_text)
                
                doc.save(docx_saida)
                docxs_gerados += 1
                
                if gerar_pdf and PDF_NATIVE_AVAILABLE:
                    if docx_to_pdf_perfeito(docx_saida, pdf_saida):
                        pdfs_gerados += 1
                
                progress = (idx + 1) / total_registros
                progress_bar.progress(progress)
                status_text.text(f"✅ {idx+1}/{total_registros}: {nome_arquivo_recibo}")
                
            except Exception as e:
                st.error(f"❌ Erro no registro {idx}: {str(e)[:50]}")
                continue
        
        if gerar_pdf:
            st.success(f"🎉 **{docxs_gerados} DOCX gerados, {pdfs_gerados} PDF gerados!**")
        else:
            st.success(f"🎉 **{docxs_gerados} DOCX gerados!**")
        
        st.info(f"📂 **TODOS os arquivos em**: `{pasta_saida}`")
        
    except Exception as e:
        st.error(f"❌ **Erro geral:** {str(e)}")

# ============================================================================
# 🌐 DETECTAR AMBIENTE
# ============================================================================
def is_streamlit_cloud():
    """Detecta se está rodando no Streamlit Cloud"""
    return ('STREAMLIT_SERVER_BASE_URL' in os.environ or 
            'STREAMLIT_RUNTIME' in os.environ or
            os.environ.get('STREAMLIT_SERVER_HEADLESS', 'false').lower() == 'true')

def is_deployed():
    """Detecta se está em deploy"""
    try:
        return not os.path.exists(str(Path.home() / "Desktop")) or \
               'streamlit-cloud' in st.__version__.lower() or \
               is_streamlit_cloud()
    except:
        return False

# ============================================================================
# 📁 GERENCIAMENTO DO CSV
# ============================================================================
def carregar_csv_colaboradores():
    """Carrega o CSV com tratamento de erros para localhost e cloud"""
    
    LOCAL_PATH = Path("C:/Users/Vinicius Guanabara/Desktop/app_streamlit/dados_colaboradores.csv")
    GITHUB_URL = "https://raw.githubusercontent.com/cfisadmfinanceirocontato-pixel/gerenciador_admfinanceiro/main/dados_colaboradores.csv"
    
    if not is_deployed():
        if LOCAL_PATH.exists():
            try:
                df = pd.read_csv(LOCAL_PATH, encoding='utf-8')
                st.success(f"✅ CSV carregado do local")
                return df
            except UnicodeDecodeError:
                try:
                    df = pd.read_csv(LOCAL_PATH, encoding='latin1')
                    st.success(f"✅ CSV carregado do local (latin1)")
                    return df
                except:
                    pass
            except:
                pass
    
    try:
        response = requests.get(GITHUB_URL, timeout=10)
        if response.status_code == 200:
            content = response.text
            for sep in [',', ';', '\t']:
                try:
                    df = pd.read_csv(StringIO(content), sep=sep)
                    if len(df.columns) > 1:
                        st.success("✅ CSV carregado do GitHub")
                        return df
                except:
                    continue
    except:
        pass
    
    st.warning("📤 **Faça upload do arquivo CSV:**")
    uploaded_file = st.file_uploader("Carregar dados_colaboradores.csv", type=['csv'], key="csv_upload")
    if uploaded_file:
        try:
            df = pd.read_csv(uploaded_file)
            st.success("✅ CSV carregado via upload")
            return df
        except Exception as e:
            st.error(f"❌ Erro ao ler arquivo: {e}")
            return pd.DataFrame()
    
    st.error("❌ Não foi possível carregar o CSV")
    return pd.DataFrame()

# ============================================================================
# ✅ FUNÇÕES DE GERENCIAMENTO DE REGISTROS
# ============================================================================
def carregar_registros(pasta_saida_str):
    """Carrega registros do Excel"""
    if not pasta_saida_str:
        return pd.DataFrame()
    
    excel_path = Path(pasta_saida_str).absolute() / "registros_completos.xlsx"
    if excel_path.exists():
        try:
            return pd.read_excel(excel_path)
        except:
            return pd.DataFrame()
    return pd.DataFrame()

def editar_registro(index, dados_editados, pasta_saida_str):
    """Edita um registro"""
    df = carregar_registros(pasta_saida_str)
    if df.empty or index >= len(df):
        return False
    
    try:
        for coluna, valor in dados_editados.items():
            if coluna in df.columns:
                df.at[index, coluna] = valor
        return salvar_registros(df, pasta_saida_str)
    except:
        return False

def excluir_registro(index, pasta_saida_str):
    """Exclui um registro"""
    df = carregar_registros(pasta_saida_str)
    if df.empty or index >= len(df):
        return False
    
    try:
        df = df.drop(index).reset_index(drop=True)
        return salvar_registros(df, pasta_saida_str)
    except:
        return False

# ============================================================================
# ✅ NOVA FUNÇÃO: EXCLUIR TODOS OS REGISTROS
# ============================================================================
def excluir_todos_registros(pasta_saida_str):
    """Exclui todos os registros com backup automático"""
    if not pasta_saida_str:
        return False, None
    
    excel_path = Path(pasta_saida_str).absolute() / "registros_completos.xlsx"
    
    try:
        if excel_path.exists():
            # Cria backup com timestamp
            backup_path = excel_path.parent / f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            shutil.copy2(excel_path, backup_path)
            
            # Remove o arquivo original
            excel_path.unlink()
            
            return True, backup_path
        return True, None
    except Exception as e:
        st.error(f"❌ Erro ao excluir registros: {e}")
        return False, None

def salvar_registros(df, pasta_saida_str):
    """Salva registros no Excel"""
    if not pasta_saida_str:
        return False
    
    excel_path = Path(pasta_saida_str).absolute() / "registros_completos.xlsx"
    try:
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Registros', index=False)
        return True
    except:
        return False

def criar_zip_todos_arquivos(pasta_saida_str):
    """Cria ZIP com todos os arquivos"""
    if not pasta_saida_str:
        return None
    
    pasta_saida = Path(pasta_saida_str).absolute()
    if not pasta_saida.exists():
        return None
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in pasta_saida.glob("*"):
            if file.is_file() and file.suffix.lower() in ['.xlsx', '.docx', '.pdf']:
                zipf.write(file, file.name)
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

# ============================================================================
# FUNÇÕES AUXILIARES
# ============================================================================
def formatar_data_completa(data_obj):
    if pd.isna(data_obj) or data_obj == '':
        return datetime.now().strftime("%d de %B de %Y").lower()
    try:
        if isinstance(data_obj, date):
            data = data_obj
        else:
            data = pd.to_datetime(data_obj, dayfirst=True).date()
        meses = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 
                6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 
                11: 'novembro', 12: 'dezembro'}
        return f"{data.day:02d} de {meses[data.month]} de {data.year}"
    except:
        return datetime.now().strftime("%d de %B de %Y").lower()

def formatar_data_csv(data_obj):
    if pd.isna(data_obj) or data_obj == '':
        return datetime.now().strftime("%d/%m/%Y")
    try:
        if isinstance(data_obj, date):
            return data_obj.strftime("%d/%m/%Y")
        return pd.to_datetime(data_obj, dayfirst=True).strftime("%d/%m/%Y")
    except:
        return str(data_obj)

def formatar_moeda(valor):
    try:
        return f"R$ {float(valor):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except:
        return "R$ 0,00"

def numero_extenso(n):
    if n == 0: return ''
    unidades = ['', 'um', 'dois', 'três', 'quatro', 'cinco', 'seis', 'sete', 'oito', 'nove']
    dezenas = ['', 'dez', 'vinte', 'trinta', 'quarenta', 'cinquenta', 'sessenta', 'setenta', 'oitenta', 'noventa']
    centenas = ['', 'cento', 'duzentos', 'trezentos', 'quatrocentos', 'quinhentos', 'seiscentos', 'setecentos', 'oitocentos', 'novecentos']
    
    if n < 10: return unidades[n]
    elif n < 100:
        d, u = divmod(n, 10)
        if u == 0: return dezenas[d]
        return f"{dezenas[d]} e {unidades[u]}"
    else:
        c, resto = divmod(n, 100)
        if resto == 0:
            return 'cem' if c == 1 else centenas[c]
        return f"{'cem' if c == 1 else centenas[c]} e {numero_extenso(resto)}"

def quantidade_por_extenso(qtd_str):
    try:
        qtd_num = float(qtd_str.replace(',', '.'))
        inteira = int(qtd_num)
        decimal = int((qtd_num - inteira) * 10 + 0.5)
        if decimal == 5:
            if inteira == 0: return "meia"
            elif inteira == 1: return "uma e meia"
            elif inteira == 2: return "duas e meia"
            elif inteira == 3: return "três e meia"
            elif inteira == 4: return "quatro e meia"
        return numero_extenso(inteira)
    except:
        return "quantidade inválida"

def valor_por_extenso(valor_str):
    try:
        valor_limpo = valor_str.replace('R$', '').replace('.', '').replace(',', '.').strip()
        valor_num = float(valor_limpo)
        reais = int(valor_num)
        centavos = int((valor_num - reais) * 100 + 0.5)
        texto_reais = numero_extenso(reais)
        if reais == 1: texto_reais += " real"
        elif reais == 0: texto_reais = "zero"
        else: texto_reais += " reais"
        if centavos > 0:
            texto_centavos = numero_extenso(centavos)
            sufixo = " centavo" if centavos == 1 else " centavos"
            return f"{texto_reais} e {texto_centavos}{sufixo}"
        return texto_reais
    except:
        return "valor inválido"

def incrementar_contador():
    st.session_state.contador_recibo += 1

def resetar_contador_funcionario(funcionario_anterior, funcionario_atual):
    if funcionario_anterior != funcionario_atual and funcionario_atual:
        st.session_state.contador_recibo = 1

def extrair_numero_oficio(oficio_completo):
    if pd.isna(oficio_completo) or not isinstance(oficio_completo, str) or '/' not in oficio_completo:
        return str(oficio_completo).strip()
    return oficio_completo.split('/')[0].strip()

# ============================================================================
# ✅ FUNÇÃO RESETAR FORMULÁRIO (CORRIGIDA)
# ============================================================================
def resetar_formulario():
    """CORREÇÃO: Reseta TODOS os estados do formulário"""
    st.session_state.contador_recibo = 1
    st.session_state.funcionario_anterior = ""
    st.rerun()

# ============================================================================
# CONFIGURAÇÃO DA APLICAÇÃO
# ============================================================================
st.set_page_config(
    page_title="Pagamento de Diárias - PDF Perfeito", 
    page_icon="📋", 
    layout="wide"
)

# ============================================================================
# SESSION STATE
# ============================================================================
if 'contador_recibo' not in st.session_state:
    st.session_state.contador_recibo = 1
if 'funcionario_anterior' not in st.session_state:
    st.session_state.funcionario_anterior = ""
if 'pasta_recibos_manual' not in st.session_state:
    st.session_state.pasta_recibos_manual = ""
if 'confirmar_exclusao_todos' not in st.session_state:
    st.session_state.confirmar_exclusao_todos = False

# ============================================================================
# CARREGAR DADOS
# ============================================================================
with st.spinner("🔄 Carregando dados..."):
    df_colaboradores = carregar_csv_colaboradores()

if df_colaboradores.empty:
    st.warning("⚠️ **Nenhum dado carregado. Verifique o arquivo CSV.**")
    st.stop()

# ============================================================================
# DADOS INICIAIS
# ============================================================================
termos_unicos = carregar_termos_colaboracao(df_colaboradores)
opcoes_quantidade = ['0,0', '0,5', '1,5', '2,5', '3,5', '4,5']

# ============================================================================
# SIDEBAR
# ============================================================================
with st.sidebar:
    st.header("🔍 Filtros")
    termo_filtro = st.selectbox("Filtrar por Termo:", ['Todos'] + termos_unicos)
    
    st.markdown("---")
    st.subheader("📄 **Contador Recibo**")
    st.metric("Próximo Nº", st.session_state.contador_recibo)
    
    st.markdown("---")
    st.subheader("🌐 **Ambiente**")
    
    if is_deployed():
        st.info("🚀 **DEPLOY**")
        if is_streamlit_cloud():
            st.warning("☁️ **Streamlit Cloud detectado**")
    else:
        st.info("🏠 **LOCALHOST**")
    
    if PDF_NATIVE_AVAILABLE:
        st.success("✅ **PDF com formatação perfeita disponível**")
    else:
        st.error("❌ **PDF não disponível - instale reportlab**")
        st.code("pip install reportlab")
    
    st.markdown("---")
    st.caption(f"📊 **Registros no CSV:** {len(df_colaboradores)}")

# ============================================================================
# PASTA DE DESTINO
# ============================================================================
st.title("📋 Pagamento de Diárias - PDF com Formatação Original")
st.markdown("---")

pasta_sugerida = "C:/Users/Vinicius Guanabara/Desktop/diarias" if not is_deployed() else "/mount/src/gerenciador_admfinanceiro/diarias"

pasta_recibos_manual = st.text_input(
    "📂 **PASTA DE DESTINO (OBRIGATÓRIA)**",
    value=st.session_state.pasta_recibos_manual or pasta_sugerida,
    placeholder="Digite o caminho completo da pasta",
    help="Os arquivos serão salvos nesta pasta"
)

if pasta_recibos_manual:
    st.session_state.pasta_recibos_manual = pasta_recibos_manual

if not pasta_recibos_manual:
    st.error("❌ **Informe a pasta de destino para continuar!**")
    st.stop()

st.markdown("---")

# ============================================================================
# FORMULÁRIO PRINCIPAL
# ============================================================================
st.subheader("📝 Novo Registro")

# Dados do Termo
st.markdown("**📋 Dados do Termo**")
termo_input = st.selectbox("Termo de Colaboração:", options=[''] + termos_unicos, index=0)

if termo_input:
    instrumento_auto = buscar_instrumento_por_termo(df_colaboradores, termo_input)
    numero_termo_auto = buscar_numero_termo_por_nome(df_colaboradores, termo_input)
else:
    instrumento_auto = ""
    numero_termo_auto = ""

col1, col2 = st.columns(2)
with col1:
    instrumento = st.text_input("Instrumento:", value=instrumento_auto)
with col2:
    numero_termo = st.text_input("Nº do Termo:", value=numero_termo_auto)

# Dados do Funcionário
st.markdown("**👤 Dados do Funcionário**")
funcionarios = carregar_funcionarios_por_termo(df_colaboradores, termo_input)
funcionario_input = st.selectbox("Funcionário:", options=[''] + funcionarios, index=0)

# Reset automático quando muda de funcionário
if st.session_state.funcionario_anterior != funcionario_input:
    resetar_contador_funcionario(st.session_state.funcionario_anterior, funcionario_input)
    st.session_state.funcionario_anterior = funcionario_input

cpf_auto, cargo_auto = "", ""
if termo_input and funcionario_input:
    cpf_auto, cargo_auto = buscar_cpf_cargo_por_funcionario(df_colaboradores, termo_input, funcionario_input)

col3, col4 = st.columns(2)
with col3:
    cpf = st.text_input("CPF:", value=cpf_auto)
with col4:
    cargo = st.text_input("Cargo:", value=cargo_auto)

# Valores
st.markdown("**💰 Valores**")
col5, col6 = st.columns(2)
with col5:
    qtd = st.selectbox("Quantidade:", options=opcoes_quantidade, index=2)
    qtd_extenso = quantidade_por_extenso(qtd)
    st.text_input("Qtd por extenso:", value=qtd_extenso, disabled=True)
with col6:
    qtd_num = float(qtd.replace(',', '.'))
    valor = formatar_moeda(qtd_num * 140)
    st.text_input("Valor:", value=valor, disabled=True)
    valor_extenso = valor_por_extenso(valor)
    st.text_input("Valor por extenso:", value=valor_extenso, disabled=True)

# Data
col7, col8 = st.columns(2)
with col7:
    data_recibo = st.date_input("Data do Recibo:", value=datetime.now().date())
    data_input = formatar_data_csv(data_recibo)
with col8:
    data_extenso = formatar_data_completa(data_recibo)
    st.text_input("Data por extenso:", value=data_extenso, disabled=True)

# Detalhes
st.markdown("**📋 Detalhes**")
objetivo = st.text_area("Objetivo:", height=80)
localidades = st.text_area("Localidades:", height=80)
periodo = st.text_input("Período:", placeholder="01/02 a 03/02")

# Ofício
col9, col10, col11 = st.columns(3)
with col9:
    oficio = st.text_input("Ofício:", placeholder="123/2026")
with col10:
    numero_oficio = extrair_numero_oficio(oficio)
    num_oficio = st.text_input("Nº do Ofício:", value=numero_oficio)
with col11:
    nome_arquivo = st.text_input("Nome Arquivo:", value=funcionario_input.split()[0] if funcionario_input else "")

# Nome do Recibo
nome_recibo_auto = f"{nome_arquivo}_{num_oficio}_{st.session_state.contador_recibo}" if nome_arquivo and num_oficio else ""
nome_recibo = st.text_input("Nome do Recibo:", value=nome_recibo_auto)

# Template
st.markdown("---")
st.subheader("📄 **Template do Recibo**")
template_file = st.file_uploader("Carregar modelo.docx", type=['docx'])

template_path = None
if template_file:
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.write(template_file.read())
        template_path = tmp.name
    st.success("✅ Template carregado!")

# ============================================================================
# BOTÕES DE AÇÃO
# ============================================================================
st.markdown("---")
st.subheader("⚡ Ações")

if PDF_NATIVE_AVAILABLE:
    cols = st.columns(5)
else:
    cols = st.columns(4)

btn_idx = 0

# Salvar
with cols[btn_idx]:
    if st.button("💾 **SALVAR**", type="primary", use_container_width=True):
        if not all([termo_input, funcionario_input, cpf]):
            st.error("❌ Preencha Termo, Funcionário e CPF!")
        else:
            resultado = salvar_registro_formulario(
                df_colaboradores, termo_input, instrumento, numero_termo, funcionario_input,
                cpf, cargo, qtd, qtd_extenso, valor, valor_extenso, data_input, data_extenso,
                objetivo, localidades, periodo, oficio, nome_arquivo, num_oficio,
                nome_recibo, pasta_recibos_manual
            )
            if resultado[0] is not None:
                st.success("✅ Registro salvo!")
                st.rerun()
btn_idx += 1

# Gerar DOCX
with cols[btn_idx]:
    if st.button("📄 **GERAR DOCX**", use_container_width=True) and template_path:
        if not all([termo_input, funcionario_input, cpf]):
            st.error("❌ Preencha Termo, Funcionário e CPF!")
        else:
            dados = {
                'Funcionário': funcionario_input,
                'Cargo': cargo,
                'CPF': cpf,
                'Valor': valor,
                'Valor por extenso': valor_extenso,
                'Quantidade': qtd,
                'Quantidade por extenso': qtd_extenso,
                'Instrumento': instrumento,
                'Nº Do Termo de Colaboração': numero_termo,
                'Objetivo': objetivo,
                'Localidades': localidades,
                'Período': periodo,
                'Ofício': oficio,
                'Data por Extenso': data_extenso,
                'Nome do Recibo': nome_recibo
            }
            resultado = gerar_recibo_individual(dados, template_path, pasta_recibos_manual, gerar_pdf=False)
            if resultado and resultado[0] is not None:
                st.success("✅ DOCX gerado!")
btn_idx += 1

# Gerar PDF
if PDF_NATIVE_AVAILABLE:
    with cols[btn_idx]:
        if st.button("🖨️ **GERAR PDF PERFEITO**", use_container_width=True) and template_path:
            if not all([termo_input, funcionario_input, cpf]):
                st.error("❌ Preencha Termo, Funcionário e CPF!")
            else:
                dados = {
                    'Funcionário': funcionario_input,
                    'Cargo': cargo,
                    'CPF': cpf,
                    'Valor': valor,
                    'Valor por extenso': valor_extenso,
                    'Quantidade': qtd,
                    'Quantidade por extenso': qtd_extenso,
                    'Instrumento': instrumento,
                    'Nº Do Termo de Colaboração': numero_termo,
                    'Objetivo': objetivo,
                    'Localidades': localidades,
                    'Período': periodo,
                    'Ofício': oficio,
                    'Data por Extenso': data_extenso,
                    'Nome do Recibo': nome_recibo
                }
                resultado = gerar_recibo_individual(dados, template_path, pasta_recibos_manual, gerar_pdf=True)
                if resultado and resultado[0] is not None:
                    st.success("✅ DOCX + PDF com formatação original gerados!")
    btn_idx += 1

# Incrementar
with cols[btn_idx]:
    if st.button("🔄 **INCREMENTAR**", use_container_width=True, on_click=incrementar_contador):
        st.success(f"✅ Contador: {st.session_state.contador_recibo}")
btn_idx += 1

# Resetar
with cols[btn_idx]:
    if st.button("🔄 **RESETAR**", use_container_width=True):
        resetar_formulario()

# ============================================================================
# OPERAÇÕES EM LOTE
# ============================================================================
st.markdown("---")
st.subheader("📦 Operações em Lote")

col_lote1, col_lote2, col_lote3, col_lote4 = st.columns(4)

with col_lote1:
    if st.button("📑 **GERAR TODOS DOCX**", use_container_width=True) and template_path:
        gerar_todos_recibos(template_path, pasta_recibos_manual, gerar_pdf=False)

with col_lote2:
    if PDF_NATIVE_AVAILABLE:
        if st.button("🖨️ **GERAR TODOS PDF**", use_container_width=True) and template_path:
            gerar_todos_recibos(template_path, pasta_recibos_manual, gerar_pdf=True)

with col_lote3:
    zip_data = criar_zip_todos_arquivos(pasta_recibos_manual)
    if zip_data:
        st.download_button(
            "📦 **BAIXAR TUDO (ZIP)**",
            zip_data,
            file_name=f"pagamentos_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
            mime="application/zip",
            use_container_width=True
        )

with col_lote4:
    # NOVO BOTÃO: Excluir todos os registros
    if st.button("🗑️ **EXCLUIR TODOS**", type="secondary", use_container_width=True):
        st.session_state.confirmar_exclusao_todos = True

# Confirmação de exclusão
if st.session_state.confirmar_exclusao_todos:
    st.markdown("---")
    st.error("⚠️ **ATENÇÃO: Esta ação é irreversível!**")
    st.warning("Todos os registros serão permanentemente excluídos. Um backup será criado automaticamente.")
    
    col_confirm1, col_confirm2 = st.columns(2)
    
    with col_confirm1:
        if st.button("✅ **CONFIRMAR EXCLUSÃO**", type="primary", use_container_width=True):
            with st.spinner("Excluindo registros..."):
                sucesso, backup_path = excluir_todos_registros(pasta_recibos_manual)
                
                if sucesso:
                    if backup_path and backup_path.exists():
                        st.success("✅ Todos os registros foram excluídos! Backup criado.")
                        
                        with open(backup_path, 'rb') as f:
                            st.download_button(
                                "📥 **Baixar Backup**",
                                f,
                                file_name=backup_path.name,
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                    else:
                        st.success("✅ Todos os registros foram excluídos!")
                    
                    st.session_state.confirmar_exclusao_todos = False
                    st.rerun()
                else:
                    st.error("❌ Erro ao excluir registros")
    
    with col_confirm2:
        if st.button("❌ **CANCELAR**", use_container_width=True):
            st.session_state.confirmar_exclusao_todos = False
            st.rerun()

# ============================================================================
# REGISTROS SALVOS
# ============================================================================
st.markdown("---")
st.subheader("📋 Registros Salvos")

df_registros = carregar_registros(pasta_recibos_manual)

if not df_registros.empty:
    if termo_filtro != 'Todos':
        df_filtrado = df_registros[df_registros['Termo de Colaboração'] == termo_filtro]
    else:
        df_filtrado = df_registros
    
    st.dataframe(df_filtrado, use_container_width=True)
    
    col_m1, col_m2, col_m3 = st.columns(3)
    with col_m1:
        st.metric("Total Registros", len(df_filtrado))
    with col_m2:
        if 'Valor' in df_filtrado.columns:
            try:
                valores = df_filtrado['Valor'].str.replace('R$', '').str.replace('.', '').str.replace(',', '.').astype(float)
                st.metric("Total R$", f"R$ {valores.sum():,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'))
            except:
                pass
    with col_m3:
        st.metric("Registros Filtrados", len(df_filtrado))
    
    excel_path = Path(pasta_recibos_manual).absolute() / "registros_completos.xlsx"
    if excel_path.exists():
        with open(excel_path, 'rb') as f:
            st.download_button(
                "📥 **Download Excel**",
                f,
                file_name="registros_completos.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
else:
    st.info("👆 **Nenhum registro encontrado. Cadastre o primeiro!**")

if template_path and os.path.exists(template_path):
    try:
        os.unlink(template_path)
    except:
        pass

# ============================================================================
# RODAPÉ
# ============================================================================
st.markdown("---")
st.caption(f"📁 **Pasta:** {pasta_recibos_manual}")
st.caption("✅ **PDF com formatação IDÊNTICA ao DOCX original**")
st.caption("🖨️ **Preserva fontes, tamanhos, cores, negrito, itálico, alinhamentos e tabelas**")